#!/usr/bin/env python3
# 3GPP 스펙 문서(.pdf/.doc/.docx)에서 텍스트를 추출해 stdout 으로 출력하는 헬퍼
"""
Usage:
    python3 design/scripts/extract.py <path-to-document> [--max-chars N]

Supported formats:
    .pdf  → pypdf
    .docx → python-docx
    .doc  → soffice (LibreOffice headless) → .docx 변환 후 python-docx
            폴백: antiword (있으면)

의존성:
    venv 에서 `pypdf`, `python-docx` 필요.
    .doc 처리에는 시스템에 `soffice` 또는 `antiword` 중 하나 필요.
"""

from __future__ import annotations

import argparse
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

DEFAULT_MAX_CHARS = 12_000


def fail(msg: str, code: int = 1) -> None:
    print(f"[extract.py] {msg}", file=sys.stderr)
    sys.exit(code)


def extract_pdf(path: Path, max_chars: int) -> str:
    try:
        import pypdf  # type: ignore
    except ImportError:
        fail(
            "pypdf 미설치. 설치: \n"
            "  python3 -m venv .venv && .venv/bin/pip install pypdf python-docx\n"
            "이후 `.venv/bin/python3 design/scripts/extract.py ...` 로 실행."
        )
    reader = pypdf.PdfReader(str(path))
    chunks: list[str] = []
    total = 0
    for page in reader.pages:
        t = page.extract_text() or ""
        if t:
            chunks.append(t)
            total += len(t)
        if total >= max_chars:
            break
    return "\n".join(chunks)[:max_chars]


def extract_docx(path: Path, max_chars: int) -> str:
    try:
        import docx  # type: ignore  # python-docx
    except ImportError:
        fail(
            "python-docx 미설치. 설치:\n"
            "  python3 -m venv .venv && .venv/bin/pip install pypdf python-docx"
        )
    doc = docx.Document(str(path))
    chunks: list[str] = []
    total = 0

    def push(text: str) -> bool:
        nonlocal total
        if not text:
            return False
        chunks.append(text)
        total += len(text)
        return total >= max_chars

    for para in doc.paragraphs:
        if push(para.text):
            return "\n".join(chunks)[:max_chars]

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if push(row_text):
                return "\n".join(chunks)[:max_chars]

    return "\n".join(chunks)[:max_chars]


def extract_doc(path: Path, max_chars: int) -> str:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice:
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            try:
                subprocess.run(
                    [
                        soffice,
                        "--headless",
                        "--convert-to",
                        "docx",
                        "--outdir",
                        str(tmp_path),
                        str(path),
                    ],
                    check=True,
                    capture_output=True,
                    timeout=120,
                )
            except subprocess.CalledProcessError as e:
                fail(
                    f"soffice 변환 실패: {e.stderr.decode(errors='replace')[:500]}"
                )
            converted = tmp_path / (path.stem + ".docx")
            if not converted.exists():
                fail(f"soffice 변환 산출물 없음: {converted}")
            return extract_docx(converted, max_chars)

    antiword = shutil.which("antiword")
    if antiword:
        try:
            out = subprocess.run(
                [antiword, str(path)],
                check=True,
                capture_output=True,
                timeout=60,
            )
            return out.stdout.decode("utf-8", errors="replace")[:max_chars]
        except subprocess.CalledProcessError as e:
            fail(f"antiword 실패: {e.stderr.decode(errors='replace')[:500]}")

    fail(
        ".doc 처리에는 soffice 또는 antiword 가 필요합니다.\n"
        "  sudo apt install libreoffice-core   # 권장 (포맷 정확도 높음)\n"
        "  sudo apt install antiword            # 가벼운 대안"
    )
    return ""


EXTRACTORS = {
    ".pdf": extract_pdf,
    ".docx": extract_docx,
    ".doc": extract_doc,
}


def main() -> None:
    parser = argparse.ArgumentParser(
        description="3GPP 스펙 문서(.pdf/.doc/.docx) 텍스트 추출"
    )
    parser.add_argument("path", type=Path, help="추출할 문서 경로")
    parser.add_argument(
        "--max-chars",
        type=int,
        default=DEFAULT_MAX_CHARS,
        help=f"출력 최대 문자 수 (기본 {DEFAULT_MAX_CHARS})",
    )
    args = parser.parse_args()

    if not args.path.exists():
        fail(f"파일 없음: {args.path}")

    suffix = args.path.suffix.lower()
    extractor = EXTRACTORS.get(suffix)
    if not extractor:
        fail(
            f"지원하지 않는 확장자: {suffix}. "
            f"지원: {', '.join(EXTRACTORS.keys())}"
        )

    text = extractor(args.path, args.max_chars)
    sys.stdout.write(text)
    if text and not text.endswith("\n"):
        sys.stdout.write("\n")


if __name__ == "__main__":
    main()
